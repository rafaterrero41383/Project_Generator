import io
import tempfile, zipfile, re, os, sys, types
import xml.etree.ElementTree as ET
import yaml
from pathlib import Path
import hashlib, shutil, time
import json
import socket
import datetime
import urllib.parse
import requests  # Para enviar logs a Datadog

# --- Parche compatibilidad 3.13 ---
if 'imghdr' not in sys.modules:
    imghdr = types.ModuleType("imghdr")
    sys.modules['imghdr'] = imghdr
# ----------------------------------

from docx import Document
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

# ========= CONFIG =========
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    st.error("❌ Falta OPENAI_API_KEY en secretos/entorno.")
    st.stop()

# Datadog (env)
DD_API_KEY = os.getenv("DATADOG_API_KEY")
DD_SITE = os.getenv("DATADOG_SITE", "datadoghq.com").strip() or "datadoghq.com"
DD_URL = f"https://http-intake.logs.{DD_SITE}/api/v2/logs"
DD_SERVICE = os.getenv("DATADOG_SERVICE", "mule-project-generator")
DD_SOURCE = os.getenv("DATADOG_SOURCE", "mulesoft-gen")
DD_TAGS = os.getenv("DATADOG_TAGS", "app:mule-gen,env:local")  # formato "k:v,k:v"

client = OpenAI()
MODEL_BASE = "gpt-4o-mini"

st.set_page_config(page_title="🤖 Generador de Proyectos", layout="wide")

# ====== UI base ======
st.markdown("""
<style>
#MainMenu { display:none }
.chat-message { display:flex; align-items:flex-start; gap:12px; margin-bottom:14px; }
.user-message { flex-direction: row-reverse; }
.avatar { width: 36px; height: 36px; border-radius: 50%; object-fit: cover; }
.message-bubble { padding: 12px 14px; border-radius: 14px; max-width: 85%; line-height:1.38; }
.user-bubble { background:#e3f2fd; border:1px solid #bbdefb; }
.assistant-bubble { background:#f1f0f0; border:1px solid #ddd; }
.sev-CRIT { color:#b71c1c; }
.sev-WARN { color:#e65100; }
.sev-INFO { color:#1565c0; }
</style>
""", unsafe_allow_html=True)

assistant_avatar = "https://cdn-icons-png.flaticon.com/512/4712/4712109.png"
user_avatar = "https://cdn-icons-png.flaticon.com/512/1077/1077012.png"

st.markdown("<h1 style='text-align:center;'>🤖 Generador de Proyectos</h1>", unsafe_allow_html=True)

# ====== Estado ======
if "messages" not in st.session_state: st.session_state.messages = []
if "uploaded_spec" not in st.session_state: st.session_state.uploaded_spec = None
if "generated_zip" not in st.session_state: st.session_state.generated_zip = None
if "observaciones" not in st.session_state: st.session_state.observaciones = []
if "service_type" not in st.session_state: st.session_state.service_type = "UNKNOWN"  # REC|DOM|BUS|PROXY|UNKNOWN
if "spec_name" not in st.session_state: st.session_state.spec_name = None
if "spec_kind" not in st.session_state: st.session_state.spec_kind = None   # "RAML" | "OAS" | "TEXT" | "ZIP"
if "is_generating" not in st.session_state: st.session_state.is_generating = False
if "pending_action" not in st.session_state: st.session_state.pending_action = None
if "archetype_choice" not in st.session_state: st.session_state.archetype_choice = "Automático"
if "rubrics_defs" not in st.session_state: st.session_state.rubrics_defs = []
if "rubrics_kind" not in st.session_state: st.session_state.rubrics_kind = "mule"  # mule|apigee
if "ctx_text" not in st.session_state: st.session_state.ctx_text = ""

TYPE_LABELS = {
    "REC": "RECEPTION",
    "DOM": "DOMAIN",
    "BUS": "BUSINESS",
    "PROXY": "PROXY",
    "UNKNOWN": "UNKNOWN"
}

col1, col2 = st.columns([1,1])
with col1:
    if st.button("🔄 Reiniciar"):
        for k in list(st.session_state.keys()): del st.session_state[k]
        st.rerun()
with col2:
    st.caption("")

# ========= Utilidades =========

TEXT_EXTS = {".xml",".json",".yaml",".yml",".raml",".properties",".txt",".pom",".md",".js",".gradle",".groovy"}
INVALID_WIN = r'[:*?"<>|\\/]'

def safe_filename(stx: str, fallback: str = "root") -> str:
    s = (stx or "").strip()
    if not s:
        return fallback
    s = re.sub(INVALID_WIN, "-", s)
    s = s.strip("-._ ")
    return s or fallback

def _map_prefix_to_type(filename: str) -> str | None:
    """Detecta tipo por prefijo/patrón en el nombre del archivo ZIP de diseño."""
    if not filename:
        return None
    fname = filename.lower()
    m = re.match(r"^(rec|dom|bus|pro)[-_]", fname)
    if m:
        pref = m.group(1)
        return {"rec":"REC","dom":"DOM","bus":"BUS","pro":"PROXY"}.get(pref)
    if "-rec-" in fname or fname.startswith("rec-") or fname.endswith("-rec.zip"):
        return "REC"
    if "-dom-" in fname or fname.startswith("dom-"):
        return "DOM"
    if "-bus-" in fname or fname.startswith("bus-"):
        return "BUS"
    if "proxy" in fname or "-pro-" in fname:
        return "PROXY"
    return None

def leer_especificacion(file) -> str:
    """
    Lee la especificación subida. Si es ZIP, extrae el mejor candidato y
    deja en session_state:
      - extracted_kind, extracted_name, extracted_bytes, ctx_text
    Devuelve texto para armar el contexto del LLM.
    """
    name = (file.name or "").lower()
    file.seek(0)

    # === ZIP de diseño ===
    if name.endswith(".zip"):
        st.session_state.spec_kind = "ZIP"
        data = file.read()
        with zipfile.ZipFile(io.BytesIO(data), "r") as z:
            kind, inner_name, inner_bytes = _best_candidate_from_zip(z)

        st.session_state.extracted_kind = kind
        st.session_state.extracted_name = inner_name
        st.session_state.extracted_bytes = inner_bytes

        if kind == "TEXT":
            ctx_text = _read_docx_bytes(inner_bytes) if inner_name.lower().endswith(".docx") \
                       else inner_bytes.decode("utf-8", "ignore")
        else:
            try:
                ctx_text = inner_bytes.decode("utf-8", "ignore")
            except Exception:
                ctx_text = ""
        st.session_state.ctx_text = ctx_text
        return ctx_text

    # === Compatibilidad archivos sueltos ===
    if name.endswith(".raml"):
        st.session_state.spec_kind = "RAML"
        return file.read().decode("utf-8", errors="ignore")

    if name.endswith((".yaml", ".yml", ".json")):
        st.session_state.spec_kind = "OAS"
        return file.read().decode("utf-8", errors="ignore")

    if name.endswith(".docx"):
        st.session_state.spec_kind = "TEXT"
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file.read())
            path = tmp.name
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    st.session_state.spec_kind = None
    return ""

def _read_docx_bytes(b: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(b)
        path = tmp.name
    try:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    finally:
        try: os.remove(path)
        except: pass

def _best_candidate_from_zip(z: zipfile.ZipFile) -> tuple[str,str,bytes]:
    """
    Retorna (kind, inner_name, bytes) priorizando:
    1) openapi.(yaml|yml|json) o *openapi*.*
    2) *.raml
    3) *.docx
    4) README(.md|.txt)
    5) cualquier .yaml/.json como último recurso
    kind ∈ {"OAS","RAML","TEXT","RAW"}
    """
    names = z.namelist()
    # 1) OpenAPI
    cand = [n for n in names if re.search(r"(^|/)(openapi\.(ya?ml|json))$", n, re.I)]
    cand += [n for n in names if re.search(r"openapi.*\.(ya?ml|json)$", n, re.I)]
    if cand:
        n = cand[0]; return ("OAS", n, z.read(n))

    # 2) RAML
    raml = [n for n in names if n.lower().endswith(".raml")]
    if raml:
        n = raml[0]; return ("RAML", n, z.read(n))

    # 3) DOCX
    docx = [n for n in names if n.lower().endswith(".docx")]
    if docx:
        n = docx[0]; return ("TEXT", n, z.read(n))

    # 4) README
    rd = [n for n in names if re.search(r"readme(\.md|\.txt)?$", n, re.I)]
    if rd:
        n = rd[0]; return ("TEXT", n, z.read(n))

    # 5) YAML/JSON genérico
    yj = [n for n in names if re.search(r"\.(ya?ml|json)$", n, re.I)]
    if yj:
        n = yj[0]
        ext = n.split(".")[-1].lower()
        kind = "OAS" if ext in ("yaml","yml","json") else "RAW"
        return (kind, n, z.read(n))

    return ("RAW", names[0] if names else "bundle.zip", z.read(names[0])) if names else ("RAW","bundle.zip",b"")

# ✅ Validadores
def validar_xml(txt: str, archivo: str) -> str|None:
    try:
        ET.fromstring(txt)
        return None
    except ET.ParseError as e:
        return f"❌ {archivo}: Error XML → {e}"

def validar_yaml(txt: str, archivo: str) -> str|None:
    try:
        yaml.safe_load(txt)
        return None
    except yaml.YAMLError as e:
        return f"⚠️ {archivo}: Error YAML → {e}"

# === Arquetipos ===

def _file_sha256(path: str, chunk=1024*1024) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            b = f.read(chunk)
            if not b: break
            h.update(b)
    return h.hexdigest()

def _prefer_zip_or_dir(candidates: list[str]) -> str | None:
    for c in candidates:
        p = Path(c)
        if p.exists(): return str(p.resolve())
    # Búsqueda heurística en cwd
    for f in os.listdir():
        if any(key in f.lower() for key in ["mx-ms-bc-rec-", "apiproxy", "apigee", "reception"]):
            return str(Path(f).resolve())
    return None

def obtener_arquetipo_generic() -> str | None:
    # Arquetipo genérico Mule (DOM/BUS)
    candidates = []
    for f in os.listdir():
        p = Path(f)
        if p.is_dir() and ("arquetipo" in p.name.lower()) and ("reception" not in p.name.lower()):
            candidates.append(str(p.resolve()))
        if p.is_file() and p.suffix.lower()==".zip" and ("arquetipo" in p.name.lower()) and ("reception" not in p.name.lower()):
            candidates.append(str(p.resolve()))
        if p.is_file() and p.suffix.lower()==".zip" and re.search(r"mx-ms-bc-dom-.*-mule", p.name.lower()):
            candidates.append(str(p.resolve()))
    demo = "/mnt/data/arquetipo-mulesoft.zip"
    if os.path.exists(demo):
        candidates.append(demo)
    return _prefer_zip_or_dir(candidates)

def obtener_arquetipo_reception() -> str | None:
    # Arquetipo REC Apigee
    candidates = []
    for f in os.listdir():
        p = Path(f)
        if p.is_dir() and ("arquetipo-reception" in p.name.lower()):
            candidates.append(str(p.resolve()))
        if p.is_file() and p.suffix.lower()==".zip" and ("arquetipo-reception" in p.name.lower()):
            candidates.append(str(p.resolve()))
        if p.is_file() and p.suffix.lower()==".zip" and re.search(r"mx-ms-bc-rec-.*apigee", p.name.lower()):
            candidates.append(str(p.resolve()))
    demo = "/mnt/data/arquetipo-reception.zip"
    if os.path.exists(demo):
        candidates.append(demo)
    return _prefer_zip_or_dir(candidates)

def preparar_arquetipo_trabajo(arquetipo_path: str) -> Path:
    p = Path(arquetipo_path)
    if p.is_dir():
        return p
    cache_root = Path.home() / ".mule_archetypes"
    cache_root.mkdir(parents=True, exist_ok=True)
    digest = _file_sha256(str(p))
    target = cache_root / digest
    marker = target / ".ready"
    if not marker.exists():
        if target.exists():
            shutil.rmtree(target, ignore_errors=True)
        target.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(str(p), "r") as z:
            z.extractall(target)
        marker.write_text(f"{int(time.time())}\n", encoding="utf-8")
    return target

# ========= Helpers PROXY =========

def _parse_base_uri(uri: str):
    if not uri: return ("http","localhost","8081","/")
    p = urllib.parse.urlparse(uri)
    protocol = p.scheme or "http"
    host = p.hostname or "localhost"
    port = str(p.port or (443 if protocol == "https" else 80))
    base_path = p.path or "/"
    if not base_path.startswith("/"): base_path = "/" + base_path
    return (protocol, host, port, base_path)

def _ensure_property_line(text: str, key: str, value: str) -> str:
    pattern = re.compile(rf"^\s*{re.escape(key)}\s*:", re.M)
    if pattern.search(text): return text
    return text.rstrip() + f"\n{key}: {value}\n"

def _patch_env_properties(root: Path, protocol: str, host: str, port: str, base_path: str):
    props_dir = root / "src" / "main" / "resources" / "properties"
    wanted = {
        "proxy.protocol": protocol,
        "proxy.host": host,
        "proxy.port": port,
        "proxy.basePath": base_path,
        "proxy.conn.timeout": "10000",
        "proxy.sock.timeout": "10000",
    }
    if not props_dir.exists(): return
    for f in props_dir.glob("*-config.yaml"):
        try:
            txt = f.read_text(encoding="utf-8")
            for k,v in wanted.items():
                txt = _ensure_property_line(txt, k, str(v))
            f.write_text(txt, encoding="utf-8")
        except Exception:
            pass

def _gen_proxy_flows(dst_mule_dir: Path, artifact_id: str):
    dst_mule_dir.mkdir(parents=True, exist_ok=True)
    xml = """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http"
      xmlns:ee="http://www.mulesoft.org/schema/mule/ee/core"
      xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
      xsi:schemaLocation="
        http://www.mulesoft.org/schema/mule/core http://www.mulesoft.org/schema/mule/core/current/mule.xsd
        http://www.mulesoft.org/schema/mule/http http://www.mulesoft.org/schema/mule/http/current/mule-http.xsd
        http://www.mulesoft.org/schema/mule/ee/core http://www.mulesoft.org/schema/mule/ee/core/current/mule-ee.xsd">

  <http:listener-config name="{ARTIFACT}-http">
    <http:listener-connection host="0.0.0.0" port="8081"/>
  </http:listener-config>

  <http:request-config name="{ARTIFACT}-upstream">
    <http:request-connection protocol="#[p('proxy.protocol')]" host="#[p('proxy.host')]" port="#[p('proxy.port')]"/>
  </http:request-config>

  <flow name="{ARTIFACT}-proxy-flow">
    <http:listener config-ref="{ARTIFACT}-http" path="/{{+proxyPath}}">
      <http:response statusCode="#[attributes.statusCode default 200]"/>
    </http:listener>

    <ee:transform doc:name="Build Target Attributes" xmlns:ee="http://www.mulesoft.org/schema/mule/ee/core">
      <ee:message>
        <ee:set-payload><![CDATA[%dw 2.0
output application/java
---
payload
]]></ee:set-payload>

        <ee:set-attributes><![CDATA[%dw 2.0
output application/java
var qp = (attributes.queryParams default {}) as Object
---
{
  method: (attributes.method default "GET") as String,
  targetPath: (p("proxy.basePath") as String) ++ "/" ++ (attributes."listenerPathParams".proxyPath as String),
  queryParams: qp,
  headers: (attributes.headers default {}) as Object
}]]></ee:set-attributes>
      </ee:message>
    </ee:transform>

    <http:request method="#[attributes.method]"
                  config-ref="{ARTIFACT}-upstream"
                  path="#[attributes.targetPath]"
                  queryParams="#[attributes.queryParams]"
                  headers="#[attributes.headers]">
      <http:request-builder>
        <http:header headerName="host" value="#[p('proxy.host')]"/>
      </http:request-builder>
    </http:request>
  </flow>

  <error-handler>
    <on-error-propagate logException="true">
      <set-payload value="#[error.description default error.message]" />
    </on-error-propagate>
  </error-handler>
</mule>
"""
    xml = xml.replace("{ARTIFACT}", artifact_id)
    (dst_mule_dir / f"{artifact_id}-proxy.xml").write_text(xml, encoding="utf-8")

# ========= PROMPT UNIFICADO (contexto + reglas de transformación) =========

PROMPT_UNIFICADO = """Responde con un ÚNICO YAML válido. Eres un generador de proyectos para cuatro capas:
- Domain (Mule 4)
- Business (Mule 4)
- Proxy (Mule 4, reverse proxy)
- Reception (Apigee)

Objetivo:
A partir de un ZIP de diseño que contenga la ESPECIFICACIÓN (OpenAPI/RAML o documento de texto) y la CAPA seleccionada ({capa}), emite un único YAML con
contexto + reglas de transformación para generar el proyecto final (sustitución de nombres, rutas y policies). El sistema extrae automáticamente
del ZIP el mejor candidato (openapi.*, *.raml o *.docx).

Entrada: un ZIP de diseño (contiene OpenAPI/RAML/TXT). Extrae el mejor candidato y úsalo como especificación fuente.
Salida: YAML unificado con CONTEXTO + TRANSFORMACIONES que el programa aplicará de forma determinista al arquetipo de la capa seleccionada.

Estructura obligatoria del YAML:

layer: domain | business | proxy | reception
names:
  project_name: string                # Mule (Domain/Business/Proxy)
  artifact_id: string-kebab           # Mule/APIGEE (como aplique)
  version: "1.0.0"
  group_id: com.company.domain        # Mule
  # Reception (Apigee):
  api_display_name: string
  api_name: string-kebab
paths:
  base_path: "/v1/resource"           # Mule listener (general.path) o Apigee <BasePath>
  base_uri: "https://host/v1"         # Mule Proxy (si aplica) o null
  target_base_url: "https://host/v1"  # Apigee TargetEndpoint
upstream:
  protocol: HTTP|HTTPS|null
  host: string|null
  path: "/v1" | "/" | null
security:
  auth: none | apikey | oauth2
  cors: true|false
  quota:
    enabled: true|false
    interval: 1
    timeUnit: minute|hour|day
    limit: 60
  spike_arrest:
    enabled: true|false
    rate: "10ps"
transformations:
  # Declarativas, se aplican al arquetipo seleccionado
  - set_mule_pom: true                          # Mule: groupId/artifactId/name/version
  - set_mule_listener_path: "${general.path}"   # Mule sin APIkit
  - set_mule_proxy_target: true                 # Mule Proxy: protocol/host/port/basePath desde base_uri
  - set_apigee_basepath: true                   # Apigee proxies/default.xml
  - set_apigee_target_url: true                 # Apigee targets/backend.xml
  - set_apigee_policies:
      cors: auto        # auto|on|off
      quota: auto
      spike_arrest: auto
      verify_apikey: auto
  - rename_api_identifiers:
      from: [ "mx-api-bc-dom-xxx", "mx-ms-bc-rec-xxx" ]
      to:   [ "${names.artifact_id}", "${names.api_name}" ]
notes: "supuestos y aclaraciones breves"

Reglas:
- Deriva artifact_id y api_name en kebab-case desde project_name/api_display_name si faltan.
- No inventes hosts/URLs si no están en la especificación: deja null y no actives la transformación asociada.
- Mantén simple: no agregues texto fuera del YAML.
"""

def _gpt(messages, temperature=0.2, model=MODEL_BASE) -> str:
    resp = client.chat.completions.create(model=model, messages=messages, temperature=temperature)
    return resp.choices[0].message.content.strip()

def inferir_yaml_unificado(contenido_api: str, layer_choice: str) -> dict:
    """Pide al LLM el YAML unificado y lo normaliza para la capa elegida."""
    layer_key = {
        "Domain": "domain",
        "Business": "business",
        "Proxy": "proxy",
        "Reception": "reception"
    }.get(layer_choice, "domain")

    yml = _gpt([
        {"role":"system","content":"Responde solo YAML válido."},
        {"role":"user","content": f"{PROMPT_UNIFICADO}\n\nCapa seleccionada: {layer_key}\n\n=== ESPECIFICACIÓN ===\n{contenido_api}"}
    ], temperature=0.1)

    m = re.search(r"```(?:yaml|yml)?\s*(.*?)```", yml, re.DOTALL)
    if m: yml = m.group(1).strip()

    try:
        data = yaml.safe_load(yml) or {}
    except Exception:
        data = {}

    # Normalización mínima
    data.setdefault("layer", layer_key)
    data.setdefault("names", {})
    data.setdefault("paths", {})
    data.setdefault("upstream", {})
    data.setdefault("security", {})
    data.setdefault("transformations", [])

    # Derivaciones
    pn = data["names"].get("project_name") or data["names"].get("api_display_name") or "MuleApplication"
    slug = re.sub(r"[^a-zA-Z0-9]+","-", pn).strip("-").lower() or "mule-application"
    data["names"].setdefault("artifact_id", slug)
    if data["layer"] == "reception":
        ad = data["names"].get("api_display_name") or pn
        apislug = re.sub(r"[^a-zA-Z0-9]+","-", ad).strip("-").lower() or "reference-data"
        data["names"].setdefault("api_display_name", ad)
        data["names"].setdefault("api_name", apislug)
        data["paths"].setdefault("base_path", "/api")
        data["paths"].setdefault("target_base_url", "https://backend.example.com")
        data["security"].setdefault("auth", "none")
        data["security"].setdefault("cors", True)
        data["security"].setdefault("quota", {"enabled": False})
        data["security"].setdefault("spike_arrest", {"enabled": False})
    else:
        data["names"].setdefault("version","1.0.0")
        data["names"].setdefault("group_id","com.company.domain" if data["layer"]=="domain" else
                                                ("com.company.business" if data["layer"]=="business" else "com.company.proxy"))
        data["paths"].setdefault("base_path", "/api/*")
    return data

# ========= Postprocesos deterministas =========

def _safe_sub(rx, text, repl_fn, count=0):
    r = re.compile(rx, re.DOTALL)
    return r.sub(lambda m: repl_fn(m), text, count=count)

def renombrar_flows_con_prefijo(xml_text: str, ctx: dict) -> str:
    artifact = ctx.get("artifact_id", "app")
    layer = ctx.get("layer_prefix")
    prefix = f"{layer}-{artifact}" if layer else artifact

    def repl(m):
        start, old, end = m.group(1), m.group(2), m.group(3)
        if old.startswith(prefix+"-") or old.startswith(artifact+"-"):
            return f'{start}{old}{end}'
        new = re.sub(r"--+", "-", f"{prefix}-{old}")
        return f'{start}{new}{end}'

    xml_text = _safe_sub(r'(<flow\s+name=")([^"]+)(")', xml_text, repl)
    xml_text = _safe_sub(r'(<sub-flow\s+name=")([^"]+)(")', xml_text, repl)
    return xml_text

def insertar_o_actualizar_tls(xml_text: str, ctx: dict) -> str:
    if not ctx.get("tls_enabled"): return xml_text
    if "xmlns:tls=" not in xml_text: return xml_text

    has_ctx = "<tls:context" in xml_text
    tls_ctx = ['<tls:context name="default-tls">']
    if ctx.get("tls_truststore_path"):
        tls_ctx.append(f'  <tls:trust-store path="{ctx["tls_truststore_path"]}" password="{ctx.get("tls_truststore_password","")}"></tls:trust-store>')
    if ctx.get("tls_keystore_path"):
        tls_ctx.append(f'  <tls:key-store path="{ctx["tls_keystore_path"]}" password="{ctx.get("tls_keystore_password","")}"></tls:key-store>')
    tls_ctx.append('</tls:context>')
    tls_block = "\n".join(tls_ctx)

    if not has_ctx:
        if "</mule>" in xml_text:
            xml_text = xml_text.replace("</mule>", tls_block + "\n</mule>")
        else:
            xml_text = xml_text + "\n" + tls_block

    def add_tls_ref(_rx: str):
        def repl(m):
            tag = m.group(0)
            if 'tlsContext-ref=' in tag: return tag
            return tag[:-1] + ' tlsContext-ref="default-tls"'
        return repl

    if (ctx.get("upstream_protocol") or "").upper() == "HTTPS" or 'protocol="HTTPS"' in xml_text:
        xml_text = _safe_sub(r'(<http:request-connection\b[^>]*)(/?>)', xml_text, add_tls_ref(r""), count=0)
        xml_text = _safe_sub(r'(<http:listener-connection\b[^>]*)(/?>)', xml_text, add_tls_ref(r""), count=0)

    return xml_text

def postprocesar_xml(xml_text: str, ctx: dict) -> str:
    xml_text = renombrar_flows_con_prefijo(xml_text, ctx)
    xml_text = insertar_o_actualizar_tls(xml_text, ctx)
    return xml_text

# ========= Parsing RAML semilight =========

HTTP_METHODS = {"get","post","put","delete","patch","head","options"}
EXTRA_ACTIONS = {"retrieve","evaluate","execute","init","create","update","delete"}

def parse_raml_semilight(raml_text: str) -> dict:
    lines = raml_text.splitlines()
    res = {}
    cur_res = None
    cur_method = None

    for i, raw in enumerate(lines):
        line = raw.rstrip()
        if not line or line.strip().startswith("#"):
            continue

        if line.lstrip().startswith("/"):
            seg = line.strip().split()[0].strip("/").split("/")[0]
            cur_res = safe_filename(seg, "root")
            res.setdefault(cur_res, {"methods": set(), "headers_required": {}, "req_types": {}, "res_types": {}})
            cur_method = None
            continue

        token = line.strip().rstrip(":").lower()
        if token in HTTP_METHODS or token in EXTRA_ACTIONS:
            cur_method = token
            res.setdefault(cur_res or "root", {"methods": set(), "headers_required": {}, "req_types": {}, "res_types": {}})
            res[cur_res or "root"]["methods"].add(cur_method)
            continue

        if cur_method and line.strip().endswith(":") and "headers" in line.lower():
            for j in range(i+1, min(i+7, len(lines))):
                l2 = lines[j].strip()
                if l2.endswith(":") and not l2.startswith("-"):
                    hdr = l2.rstrip(":")
                    block = "\n".join(lines[j: min(j+6, len(lines))]).lower()
                    if "required: true" in block:
                        res[cur_res or "root"]["headers_required"].setdefault(cur_method, []).append(hdr)

        if cur_method and "body:" in line.lower():
            block = "\n".join(lines[i: min(i+10, len(lines))])
            m = re.search(r"type:\s*([A-Za-z0-9_\-\.]+)", block)
            if m: res[cur_res or "root"]["req_types"][cur_method] = m.group(1)

        if cur_method and "responses:" in line.lower():
            block = "\n".join(lines[i: min(i+15, len(lines))])
            m = re.search(r"200:\s*(?:\n|\r\n).*?type:\s*([A-Za-z0-9_\-\.]+)", block, re.DOTALL)
            if m: res[cur_res or "root"]["res_types"][cur_method] = m.group(1)

    for r, d in res.items():
        if not d["methods"]:
            d["methods"].add("retrieve")
    return res

# ========= Generadores XML Mule =========

def _xml_header(apikit: bool=False):
    if apikit:
        return """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns:ee="http://www.mulesoft.org/schema/mule/ee/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http"
      xmlns:apikit="http://www.mulesoft.org/schema/mule/apikit"
      xmlns:tls="http://www.mulesoft.org/schema/mule/tls"
      xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
      xsi:schemaLocation="
        http://www.mulesoft.org/schema/mule/core http://www.mulesoft.org/schema/mule/core/current/mule.xsd
        http://www.mulesoft.org/schema/mule/http http://www.mulesoft.org/schema/mule/http/current/mule-http.xsd
        http://www.mulesoft.org/schema/mule/apikit http://www.mulesoft.org/schema/mule/apikit/current/mule-apikit.xsd
        http://www.mulesoft.org/schema/mule/tls http://www.mulesoft.org/schema/mule/tls/current/mule-tls.xsd
        http://www.mulesoft.org/schema/mule/ee/core http://www.mulesoft.org/schema/mule/ee/core/current/mule-ee.xsd">"""
    else:
        return """<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns:ee="http://www.mulesoft.org/schema/mule/ee/core"
      xmlns:http="http://www.mulesoft.org/schema/mule/http"
      xmlns:tls="http://www.mulesoft.org/schema/mule/tls"
      xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
      xsi:schemaLocation="
        http://www.mulesoft.org/schema/mule/core http://www.mulesoft.org/schema/mule/core/current/mule.xsd
        http://www.mulesoft.org/schema/mule/http http://www.mulesoft.org/schema/mule/http/current/mule-http.xsd
        http://www.mulesoft.org/schema/mule/tls http://www.mulesoft.org/schema/mule/tls/current/mule-tls.xsd
        http://www.mulesoft.org/schema/mule/ee/core http://www.mulesoft.org/schema/mule/ee/core/current/mule-ee.xsd">"""

def _xml_footer():
    return "</mule>\n"

def _mk_error_handler_ref():
    return """
    <error-handler>
      <on-error-continue logException="true" type="ANY">
        <flow-ref name="hdl_commonErrorHandler"/>
      </on-error-continue>
    </error-handler>""".rstrip()

def client_file_xml(resource: str, methods: set, general_path: str, use_apikit: bool, raml_cp: str|None):
    header = _xml_header(apikit=use_apikit)

    flows = []
    for m in sorted(methods):
        fname = f"{resource}_client_{m}"
        flows.append(f'''  <flow name="{fname}">
    <logger level="INFO" message="[{resource}] client {m} start"/>
    <flow-ref name="{resource}_handler_{m}"/>
    <logger level="INFO" message="[{resource}] client {m} end"/>
{_mk_error_handler_ref()}
  </flow>''')

    apikit_part = ""
    if use_apikit and raml_cp:
        apikit_part = f"""
  <http:listener-config name="{resource}-httpListener">
    <http:listener-connection host="0.0.0.0" port="${{http.port}}"/>
  </http:listener-config>

  <apikit:config name="{resource}-api-config" raml="{raml_cp}"/>

  <flow name="{resource}-http-inbound">
    <http:listener config-ref="{resource}-httpListener" path="${{general.path}}"/>
    <apikit:router config-ref="{resource}-api-config"/>
  </flow>

  <flow name="{resource}-api-console">
    <http:listener config-ref="{resource}-httpListener" path="/console/*"/>
    <apikit:console config-ref="{resource}-api-config"/>
  </flow>
""".rstrip()

        impl = []
        for m in sorted(methods):
            apiflow = f"{m}:/"+resource
            impl.append(f'''  <flow name="{apiflow}">
    <flow-ref name="{resource}_client_{m}"/>
{_mk_error_handler_ref()}
  </flow>''')
        apikit_part = apikit_part + "\n" + "\n".join(impl)

    content = f"""{header}
  <!-- client consolidado para recurso '{resource}' -->
{apikit_part if apikit_part else ""}
{chr(10).join(flows)}
{_xml_footer()}"""
    return content.strip()

def handler_file_xml(resource: str, methods: set, raml_info: dict):
    header = _xml_header(False)
    blocks = []
    for m in sorted(methods):
        req_headers = raml_info.get(resource, {}).get("headers_required", {}).get(m, []) if raml_info else []
        hdr_sets = "\n".join([f'    <set-variable variableName="{h.replace("-", "_")}" value="#[attributes.headers.\'{h}\']"/>' for h in req_headers]) if req_headers else "    <!-- no required headers detected -->"
        blocks.append(f'''  <flow name="{resource}_handler_{m}">
{hdr_sets}
    <flow-ref name="{resource}_orchestrator_{m}"/>
{_mk_error_handler_ref()}
  </flow>''')
    return f"""{header}
  <!-- handler consolidado para recurso '{resource}' -->
{chr(10).join(blocks)}
{_xml_footer()}""".strip()

def orchestrator_file_xml(resource: str, methods: set, raml_info: dict, single_file: bool):
    header = _xml_header(False)
    if single_file:
        blocks = []
        for m in sorted(methods):
            blocks.append(f'''  <flow name="{resource}_orchestrator_{m}">
    <logger level="INFO" message="[{resource}] orchestrator {m} start - reqId=#[attributes.headers.'consumerRequestId'] sess=#[attributes.headers.'sessionId']"/>
    <set-payload value='{{"status":"OK","resource":"{resource}","method":"{m}"}}' mimeType="application/json"/>
    <logger level="INFO" message="[{resource}] orchestrator {m} end"/>
{_mk_error_handler_ref()}
  </flow>''')
        return f"""{header}
  <!-- orchestrator consolidado para recurso '{resource}' -->
{chr(10).join(blocks)}
{_xml_footer()}""".strip()
    else:
        m = next(iter(methods))
        return f"""{header}
  <flow name="{resource}_orchestrator_{m}">
    <logger level="INFO" message="[{resource}] orchestrator {m} start - reqId=#[attributes.headers.'consumerRequestId'] sess=#[attributes.headers.'sessionId']"/>
    <set-payload value='{{"status":"OK","resource":"{resource}","method":"{m}"}}' mimeType="application/json"/>
    <logger level="INFO" message="[{resource}] orchestrator {m} end"/>
{_mk_error_handler_ref()}
  </flow>
{_xml_footer()}""".strip()

def common_error_handler_xml():
    header = _xml_header(False)
    return f"""{header}
  <!-- common error handler -->
  <sub-flow name="hdl_commonErrorHandler">
    <logger level="ERROR" message="type=#[error.errorType] desc=#[error.description]"/>
  </sub-flow>
{_xml_footer()}""".strip()

def common_global_config_xml(use_apikit: bool):
    header = _xml_header(False)
    listener = "" if use_apikit else """
  <http:listener-config name="global-httpListener">
    <http:listener-connection host="0.0.0.0" port="${http.port}"/>
  </http:listener-config>""".rstrip()
    return f"""{header}
  <configuration-properties file="properties/application.properties"/>
{listener}
{_xml_footer()}""".strip()

# ========= Archivos base Mule =========

def ensure_dirs(root: Path):
    base = root / "src/main/mule"
    for d in ["client", "handler", "orchestrator", "common"]:
        (base / d).mkdir(parents=True, exist_ok=True)
    (root / "src/main/resources/api").mkdir(parents=True, exist_ok=True)
    (root / "src/main/resources/properties").mkdir(parents=True, exist_ok=True)
    (root / "scripts").mkdir(parents=True, exist_ok=True)
    (root / "src/test/munit").mkdir(parents=True, exist_ok=True)

def write_minimum_base_files(root: Path):
    props = root / "src/main/resources/properties/application.properties"
    if not props.exists():
        props.write_text("http.port=8081\ngeneral.path=/api/*\n", encoding="utf-8")
    maf = root / "mule-artifact.json"
    if not maf.exists():
        maf.write_text('{"minMuleVersion":"4.6.0","secureProperties":[]}\n', encoding="utf-8")
    pom = root / "pom.xml"
    if not pom.exists():
        pom.write_text("""<project xmlns="http://maven.apache.org/POM/4.0.0" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" xsi:schemaLocation="http://maven.apache.org/POM/4.0.0 https://maven.apache.org/xsd/maven-4.0.0.xsd">
  <modelVersion>4.0.0</modelVersion>
  <groupId>com.company.experience</groupId>
  <artifactId>mule-app</artifactId>
  <version>1.0.0</version>
  <packaging>mule-application</packaging>
  <name>mule-app</name>
</project>
""", encoding="utf-8")

def first_raml_target(dst_root: Path) -> Path:
    return dst_root / "src/main/resources/api/api.raml"

def oas_target(dst_root: Path) -> Path:
    return dst_root / "src/main/resources/api/openapi.yaml"

def raml_classpath(root: Path) -> str|None:
    target = first_raml_target(root)
    try:
        rel = target.relative_to(root / "src/main/resources")
        return "classpath:/" + "/".join(rel.parts)
    except Exception:
        return None

# ========= Transformador con GPT (genérico) =========

def transformar_archivo_con_gpt(fname: str, original: str, ctx: dict) -> str:
    PROMPT_FILE = """Eres un configurador experto de proyectos Mule 4 / Apigee.
Actualiza el archivo indicado usando METADATOS (YAML unificado: contexto + transformaciones).

Reglas:
- Mantén formato/saltos.
- No agregues explicaciones ni ``` .
- Aplica solo transformaciones coherentes con el YAML (si una clave es null o falta, no inventes).
- Sustituye placeholders relevantes:
  * Mule: groupId, artifactId, version, project.mule.name, listener path/port, http-request host/protocol/path,
          properties YAML, exchange.json, nombres de flows/archivos si procede.
  * Apigee: nombre de apiproxy/bundle, BasePath, TargetEndpoint URL,
            y políticas (VerifyAPIKey/SpikeArrest/Quota/CORS) según 'security' y 'transformations'.
- No borres secciones no relacionadas.

=== METADATOS (YAML UNIFICADO) ===
{ctx_yaml}

=== ARCHIVO ({fname}) ORIGINAL ===
{original}
"""
    ctx_yaml = yaml.safe_dump(ctx, sort_keys=False, allow_unicode=True)
    raw = _gpt([
        {"role":"system","content":"Actúa como refactorizador determinista de archivos Mule/Apigee/XML/YAML/JSON/JS/properties."},
        {"role":"user","content": PROMPT_FILE.format(ctx_yaml=ctx_yaml, fname=fname, original=original)}
    ], temperature=0.1)

    blocks = re.findall(r"```(?:xml|yaml|yml|json|properties|txt|js)?\s*(.*?)```", raw, re.DOTALL)
    contenido = (blocks[-1].strip() if blocks else raw.strip())
    if not contenido or len(contenido.splitlines()) < max(3, int(len(original.splitlines())*0.3)):
        return original
    return contenido

# ========= POM ajustes =========

def enforce_pom_requirements(root_dir: Path, ctx: dict, use_apikit: bool):
    pom_path = root_dir / "pom.xml"
    if not pom_path.exists(): return
    txt = pom_path.read_text(encoding="utf-8", errors="ignore")
    try:
        tree = ET.fromstring(txt)
    except ET.ParseError:
        return
    ns = tree.tag.split("}",1)[0][1:] if tree.tag.startswith("{") else None
    def q(tag): return f"{{{ns}}}{tag}" if ns else tag

    packaging = tree.find(q("packaging"))
    if packaging is None or (packaging.text or "").strip().lower() != "mule-application":
        if packaging is None: packaging = ET.SubElement(tree, q("packaging"))
        packaging.text = "mule-application"

    artifactId = tree.find(q("artifactId"))
    name = tree.find(q("name"))
    artifact = ctx.get("artifact_id") or "mule-app"
    pname = ctx.get("project_name") or artifact
    if artifactId is None:
        artifactId = ET.SubElement(tree, q("artifactId"))
    artifactId.text = artifact
    if name is None:
        name = ET.SubElement(tree, q("name"))
    name.text = pname

    build = tree.find(q("build")) or ET.SubElement(tree, q("build"))
    plugins = build.find(q("plugins")) or ET.SubElement(build, q("plugins"))
    has_plugin = False
    for p in plugins.findall(q("plugin")):
        gid = (p.find(q("groupId")).text if p.find(q("groupId")) is not None else "")
        aid = (p.find(q("artifactId")).text if p.find(q("artifactId")) is not None else "")
        if gid.strip()=="org.mule.tools.maven" and aid.strip()=="mule-maven-plugin":
            has_plugin = True
            ext = p.find(q("extensions")) or ET.SubElement(p, q("extensions"))
            ext.text = "true"
    if not has_plugin:
        p = ET.SubElement(plugins, q("plugin"))
        ET.SubElement(p, q("groupId")).text = "org.mule.tools.maven"
        ET.SubElement(p, q("artifactId")).text = "mule-maven-plugin"
        ET.SubElement(p, q("version")).text = "4.2.0"
        ET.SubElement(p, q("extensions")).text = "true"

    if use_apikit:
        deps = tree.find(q("dependencies")) or ET.SubElement(tree, q("dependencies"))
        found = False
        for d in deps.findall(q("dependency")):
            gid = (d.find(q("groupId")).text if d.find(q("groupId")) is not None else "")
            aid = (d.find(q("artifactId")).text if d.find(q("artifactId")) is not None else "")
            if gid.strip()=="org.mule.modules" and aid.strip()=="mule-apikit-module":
                found = True
        if not found:
            d = ET.SubElement(deps, q("dependency"))
            ET.SubElement(d, q("groupId")).text = "org.mule.modules"
            ET.SubElement(d, q("artifactId")).text = "mule-apikit-module"
            ET.SubElement(d, q("version")).text = "1.11.0"
            ET.SubElement(d, q("classifier")).text = "mule-plugin"

# ========= README / scripts (Mule) =========

def write_readme(root: Path, raml_info: dict, ctx: dict, service_type: str, spec_kind: str | None):
    readme = root / "README.md"
    lines = [
        "# Proyecto generado",
        "",
        f"**Capa detectada/seleccionada:** `{TYPE_LABELS.get(service_type, service_type)}`",
        f"**Especificación:** `{spec_kind or 'N/A'}`",
        "",
        "## Árbol de carpetas (Mule)",
        "",
        "```\nsrc/main/mule/\n  client/\n  handler/\n  orchestrator/\n  common/\n```",
        "",
        "## Cómo ejecutar",
        "- Importa en Studio o empaqueta con Maven.",
        "- Puerto: `application.properties -> http.port`.",
        "- Path base: `${general.path}`.",
        "",
        "## Recursos y operaciones (si RAML)",
        "",
        "| Recurso | Operaciones | Headers requeridos |",
        "|---|---|---|",
    ]
    for r, d in sorted(raml_info.items()):
        ops = ", ".join(sorted(d.get("methods", [])))
        hdrs = {m: d.get("headers_required", {}).get(m, []) for m in d.get("methods", [])}
        hdrs_str = "; ".join([f"{m}: {', '.join(v) if v else '-'}" for m, v in hdrs.items()]) if hdrs else "-"
        lines.append(f"| {r} | {ops} | {hdrs_str} |")

    if service_type == "PROXY":
        lines += [
            "",
            "## Configuración PROXY",
            "",
            "En `src/main/resources/properties/*-config.yaml`:",
            "```yaml",
            "proxy.protocol: https",
            "proxy.host: api.backend.com",
            "proxy.port: 8443",
            "proxy.basePath: /v1",
            "proxy.conn.timeout: 10000",
            "proxy.sock.timeout: 10000",
            "```",
            "",
            "Flujo generado: `<artifactId>-proxy.xml` (listener 8081 → upstream)."
        ]
    if service_type == "REC":
        lines += [
            "",
            "## Layout Reception / Apigee",
            "La OAS se copió a `src/main/apigee/apiproxies/<bundle>/apiproxy/resources/oas/rec-reference-data.json`.",
            "Se ajustaron proxies, targets y políticas según metadatos."
        ]
    readme.write_text("\n".join(lines)+"\n", encoding="utf-8")

def write_validate_script(root: Path):
    scripts = root / "scripts"
    scripts.mkdir(parents=True, exist_ok=True)
    content = r"""#!/usr/bin/env bash
set -euo pipefail

YELLOW='\033[1;33m'; NC='\033[0m'
warn() { echo -e "${YELLOW}WARNING:${NC} $*"; }

# Solo avisos, no se corta el pipeline
for d in client handler orchestrator common; do
  test -d "src/main/mule/$d" || warn "Falta carpeta $d (recomendado por rúbricas)"
done

test -f "pom.xml" || warn "Falta pom.xml"
test -f "mule-artifact.json" || warn "Falta mule-artifact.json"
test -f "src/main/resources/properties/application.properties" || warn "Falta application.properties"

warn "Validación de estructura completada (no bloqueante)."
"""
    (scripts / "validate-structure.sh").write_text(content, encoding="utf-8")
    os.chmod(scripts / "validate-structure.sh", 0o755)

def write_munit_min(root: Path, raml_info: dict):
    tests_dir = root / "src/test/munit"
    for recurso, d in sorted(raml_info.items()):
        for m in sorted(d.get("methods", [])):
            name = f"{recurso}-{m}-test.xml"
            xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<mule xmlns:munit="http://www.mulesoft.org/schema/mule/munit"
      xmlns="http://www.mulesoft.org/schema/mule/core"
      xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
      xsi:schemaLocation="
        http://www.mulesoft.org/schema/mule/core http://www.mulesoft.org/schema/mule/core/current/mule.xsd
        http://www.mulesoft.org/schema/mule/munit http://www.mulesoft.org/schema/mule/munit/current/mule-munit.xsd">

  <munit:test name="{recurso}-{m}-happy">
    <munit:execution>
      <flow-ref name="{recurso}_client_{m}"/>
    </munit:execution>
    <munit:validation>
      <munit-tools:assert-that xmlns:munit-tools="http://www.mulesoft.org/schema/mule/munit-tools"
        expression="#[payload]" is="#[MunitTools::notNullValue()]"/>
    </munit:validation>
  </munit:test>
</mule>
"""
            (tests_dir / name).write_text(xml, encoding="utf-8")

# ========= RÚBRICAS =========

def _normalize_rubric_item(item: dict) -> dict:
    return {
        "id": item.get("id") or item.get("code") or item.get("rule_id") or "",
        "label": item.get("label") or item.get("title") or "",
        "category": item.get("category") or item.get("group") or "",
        "severity": (item.get("severity") or "WARN").upper(),
        "autofix": bool(item.get("autofix", False)),
        "action": (item.get("check", {}) or {}).get("action", "report_warning"),
        "impact": item.get("impact") or "",
        "enabled": item.get("enabled", True),
    }

def _load_rubrics_from_root(filename: str, kind_label: str) -> list[dict]:
    path = Path(filename)
    if not path.exists():
        st.sidebar.warning(f"⚠️ No se encontró {filename} en la raíz. Se aplicarán solo validaciones básicas ({kind_label}).")
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        arr = data.get("rubrics", data)
        if isinstance(arr, dict): arr = [arr]
        rubrics = [_normalize_rubric_item(x) for x in arr if isinstance(x, dict)]
        st.sidebar.success(f"✅ Rúbricas ({kind_label}) cargadas automáticamente ({len(rubrics)} definiciones).")
        return rubrics
    except Exception as e:
        st.sidebar.error(f"❌ Error al cargar rúbricas {filename}: {e}")
        return []

# Carga inicial (por defecto asume Mule)
if not st.session_state.rubrics_defs:
    st.session_state.rubrics_defs = _load_rubrics_from_root("Rubrics_Generation_Mule.json", "Mule")
    st.session_state.rubrics_kind = "mule"

def apply_rubrics_observations(root: Path, use_apikit: bool, raml_info: dict, base_notes: list[str]) -> list[str]:
    items = [r for r in st.session_state.rubrics_defs if r.get("enabled", True)]
    keywords_map = [
        ("carpeta", ["Estructura","carpeta","folder","estructura"]),
        ("XML sueltos", ["XML sueltos","suelto","mover a subcarpetas"]),
        ("pom.xml", ["pom.xml"]),
        ("mule-artifact.json", ["mule-artifact.json"]),
        ("application.properties", ["application.properties"]),
        ("listener", ["listener-config","global-httpListener","Listener"]),
        ("error handler", ["common-error-handler","hdl_commonErrorHandler","error handler"]),
        ("naming", ["Naming","-client.xml","-handler.xml","-orchestrator.xml"]),
        ("apiproxy", ["apiproxy","proxies","targets","policies"]),
        ("oas", ["openapi","oas","rec-reference-data.json"]),
        ("quota", ["Quota","cuota"]),
        ("spike", ["SpikeArrest","Spike-Arrest","spike-arrest"]),
        ("cors", ["CORS","cors"]),
        ("apikey", ["verifyapikey","apikey","api key"]),
    ]

    def guess_rubric_for_note(note: str) -> dict | None:
        note_l = note.lower()
        candidates = []
        for r in items:
            text = f"{r.get('label','')} {r.get('category','')} {r.get('impact','')}".lower()
            score = 0
            for _, keys in keywords_map:
                if any(k.lower() in note_l for k in keys):
                    if any(k.lower() in text for k in keys):
                        score += 1
            if score>0:
                candidates.append((score, r))
        if not candidates:
            return None
        candidates.sort(key=lambda x: (-x[0], x[1].get("severity","WARN")))
        return candidates[0][1]

    adorned = []
    for note in base_notes:
        rub = guess_rubric_for_note(note)
        if rub:
            sev = (rub.get("severity") or "WARN").upper()
            if sev not in ("CRIT","WARN","INFO"): sev = "WARN"
            rid = rub.get("id") or "R?"
            adorned.append(f"<span class='sev-{sev}'><b>[{sev}]</b></span> ({rid}) {note}")
        else:
            sev = "WARN"
            if re.search(r"\b(no bloqueante|recomendad[oa])\b", note, re.I):
                sev = "INFO"
            adorned.append(f"<span class='sev-{sev}'><b>[{sev}]</b></span> {note}")

    shown_ids = set()
    for line in adorned:
        m = re.search(r"\(([^)]+)\)\s", line)
        if m: shown_ids.add(m.group(1).strip())

    pendents = []
    for r in items:
        if not r.get("enabled", True):
            continue
        sev = (r.get("severity") or "WARN").upper()
        if sev not in ("CRIT","WARN"):
            continue
        rid = r.get("id") or ""
        if rid and rid in shown_ids:
            continue
        if r.get("label"):
            pendents.append(f"<span class='sev-{sev}'><b>[{sev}]</b></span> ({rid}) [Rúbrica] {r.get('label')} — pendiente de verificación automática")

    return adorned + pendents[:12]

# ---- Observaciones Mule básicas
def rubric_observaciones_basic_mule(root: Path, use_apikit: bool, raml_info: dict):
    notes = []
    base = root / "src/main/mule"
    required_dirs = ["client","handler","orchestrator","common"]
    for d in required_dirs:
        if not (base/d).exists():
            notes.append(f"[Estructura] Falta carpeta src/main/mule/{d}/")
    loose = [str(p.relative_to(root)) for p in base.glob("*.xml")]
    if loose:
        notes.append("[Estructura] XML sueltos en src/main/mule (mover a subcarpetas): " + ", ".join(loose))

    if not (root/"pom.xml").exists(): notes.append("[Activos] Falta pom.xml")
    if not (root/"mule-artifact.json").exists(): notes.append("[Activos] Falta mule-artifact.json")
    if not (root/"src/main/resources/properties/application.properties").exists(): notes.append("[Activos] Falta application.properties")

    common_txt = (base/"common"/"global-config.xml").read_text("utf-8","ignore") if (base/"common"/"global-config.xml").exists() else ""
    if use_apikit:
        if "http:listener-config" in common_txt:
            notes.append("[APIkit] Con APIkit no se recomienda listener-config en common/global-config.xml")
    else:
        if "http:listener-config" not in common_txt:
            notes.append("[Listener] Sin APIkit, conviene listener-config en common/global-config.xml")
        for folder in ["client","handler","orchestrator"]:
            for p in (base/folder).glob("*.xml"):
                if "http:listener-config" in p.read_text("utf-8","ignore"):
                    notes.append(f"[Listener] {p.name}: listener-config centralizar en common/global-config.xml")

    ceh = base/"handler"/"common-error-handler.xml"
    if not ceh.exists():
        notes.append("[Errores] Falta common-error-handler.xml")
    else:
        if "hdl_commonErrorHandler" not in ceh.read_text("utf-8","ignore"):
            notes.append("[Errores] common-error-handler.xml: definir sub-flow 'hdl_commonErrorHandler'")

    def bad_names(dirpath: Path, pattern: str, extra_ok=None):
        extra_ok = extra_ok or set()
        bad = []
        for p in dirpath.glob("*.xml"):
            if p.name in extra_ok: continue
            if not re.match(pattern, p.name): bad.append(p.name)
        return bad
    bad_client = bad_names(base/"client", r'^[a-z][A-Za-z0-9]*-client\.xml$')
    if bad_client: notes.append("[Naming] client/: " + ", ".join(bad_client))
    bad_handler = bad_names(base/"handler", r'^[a-z][A-Za-z0-9]*-handler\.xml$', {"common-error-handler.xml"})
    if bad_handler: notes.append("[Naming] handler/: " + ", ".join(bad_handler))
    bad_orch = bad_names(base/"orchestrator", r'^([a-z][A-Za-z0-9]*-(get|post|put|delete|patch|head|options|retrieve|evaluate|execute|init|create|update|delete)-orchestrator|[a-z][A-Za-z0-9]*-orchestrator)\.xml$')
    if bad_orch: notes.append("[Naming] orchestrator/: " + ", ".join(bad_orch))

    return notes

# ---- Observaciones Apigee básicas
def rubric_observaciones_basic_apigee(root: Path):
    notes = []
    base = root / "src/main/apigee/apiproxies"
    # Buscar apiproxy
    apiproxy_dirs = list(base.glob("*/apiproxy"))
    if not apiproxy_dirs:
        notes.append("[Apigee] No se encontró apiproxy bajo src/main/apigee/apiproxies/*/apiproxy")
        return notes
    apidir = apiproxy_dirs[0]
    if not (apidir/"proxies/default.xml").exists():
        notes.append("[Apigee] Falta proxies/default.xml")
    if not (apidir/"targets/backend.xml").exists():
        notes.append("[Apigee] Falta targets/backend.xml")
    if not (apidir/"resources/oas").exists():
        notes.append("[Apigee] Falta carpeta resources/oas con OAS")
    pols = list((apidir/"policies").glob("*.xml"))
    if not pols:
        notes.append("[Apigee] No hay políticas en policies/*.xml (se esperan VerifyAPIKey/SpikeArrest/Quota/AssignMessage/etc)")
    bundle_name = apidir.parent.name
    if not re.match(r"^[A-Za-z0-9._-]+$", bundle_name):
        notes.append(f"[Naming] Nombre de bundle apiproxy '{bundle_name}' contiene caracteres no recomendados")
    return notes

def rubric_observaciones(root: Path, use_apikit: bool, raml_info: dict, mode: str):
    if mode == "apigee":
        base_notes = rubric_observaciones_basic_apigee(root)
    else:
        base_notes = rubric_observaciones_basic_mule(root, use_apikit, raml_info)
    if st.session_state.rubrics_defs:
        return apply_rubrics_observations(root, use_apikit, raml_info, base_notes)
    return base_notes

# ========= LOGGING A DATADOG =========

def _backup_log_locally(payload: dict):
    try:
        logdir = Path("logs")
        logdir.mkdir(parents=True, exist_ok=True)
        ts = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        (logdir / f"chatlog_{ts}.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass  # no bloqueante

def send_chatlog_to_datadog(event: dict):
    _backup_log_locally(event)
    if not DD_API_KEY:
        st.sidebar.warning("⚠️ DATADOG_API_KEY no configurada. Log guardado localmente en ./logs/")
        return

    headers = {"Content-Type": "application/json", "DD-API-KEY": DD_API_KEY}
    body = [{
        "ddsource": DD_SOURCE,
        "service": DD_SERVICE,
        "host": socket.gethostname(),
        "ddtags": DD_TAGS,
        "message": "mule_project_generation",
        "status": "info",
        "timestamp": int(time.time() * 1000),
        "attributes": event
    }]
    try:
        r = requests.post(DD_URL, headers=headers, data=json.dumps(body), timeout=10)
        if r.status_code >= 300:
            st.sidebar.error(f"Datadog log falló ({r.status_code}): {r.text[:200]}")
        else:
            st.sidebar.info("📤 Log de generación enviado a Datadog.")
    except Exception as e:
        st.sidebar.error(f"Datadog log error: {e}")

# ========= PROCESOS DE GENERACIÓN =========

def procesar_arquetipo_mule(arquetipo_dir: str, ctx: dict, spec_bytes: bytes|None, spec_kind: str):
    src = Path(arquetipo_dir)
    tmp_dir = Path(tempfile.mkdtemp())
    shutil.copytree(src, tmp_dir, dirs_exist_ok=True)
    roots = [p for p in tmp_dir.iterdir() if p.is_dir()]
    root = roots[0] if len(roots)==1 else tmp_dir

    ensure_dirs(root)
    write_minimum_base_files(root)

    files_to_touch = []
    for r,_,fs in os.walk(root):
        for f in fs:
            p = Path(r)/f
            if p.suffix.lower() in (".png",".jpg",".jpeg",".gif",".webp",".svg",".pdf",".ppt",".pptx",".key",".ai",".psd"):
                continue
            files_to_touch.append(p)

    prog = st.progress(0.0)
    total = len(files_to_touch)
    for i, path in enumerate(files_to_touch, 1):
        prog.progress(i/total)
        try:
            if path.suffix.lower() in TEXT_EXTS or path.name.lower()=="pom.xml":
                original = path.read_text(encoding="utf-8", errors="ignore")
                nuevo = transformar_archivo_con_gpt(path.name, original, ctx)
                if path.suffix.lower() == ".xml":
                    nuevo = postprocesar_xml(nuevo, ctx)
                ext = path.suffix.lower()
                if path.name.lower()=="pom.xml" or ext==".xml":
                    err = validar_xml(nuevo, path.name)
                    if err: nuevo = original
                elif ext in (".yaml",".yml"):
                    err = validar_yaml(nuevo, path.name)
                    if err: nuevo = original
                path.write_text(nuevo, encoding="utf-8")
        except Exception:
            pass

    raml_info = {}
    raml_cp_value = None
    service_type = st.session_state.get("service_type","UNKNOWN")

    if spec_bytes:
        if spec_kind == "RAML":
            target = first_raml_target(root)
            target.parent.mkdir(parents=True, exist_ok=True)
            with open(target, "wb") as f:
                f.write(spec_bytes)
            try:
                raml_text = spec_bytes.decode("utf-8","ignore")
                raml_info = parse_raml_semilight(raml_text)
            except Exception:
                raml_info = {}
            if service_type not in ("PROXY","REC"):
                raml_cp_value = raml_classpath(root)
        elif spec_kind == "OAS":
            target = oas_target(root)
            target.parent.mkdir(parents=True, exist_ok=True)
            with open(target, "wb") as f:
                f.write(spec_bytes)

    base = root / "src/main/mule"
    client_dir = base / "client"
    handler_dir = base / "handler"
    orch_dir = base / "orchestrator"
    common_dir = base / "common"

    ceh = handler_dir / "common-error-handler.xml"
    if not ceh.exists():
        ceh.write_text(common_error_handler_xml(), encoding="utf-8")

    if service_type == "PROXY":
        (common_dir / "global-config.xml").write_text(common_global_config_xml(False), encoding="utf-8")
        _gen_proxy_flows(base, ctx.get("artifact_id","mule-app"))
        proto, host, port, bpath = _parse_base_uri((ctx.get("base_uri") or "") or "")
        _patch_env_properties(root, proto, host, port, bpath)
    else:
        use_apikit = False if service_type in ("PROXY","REC") else bool(raml_cp_value)
        gc = common_dir / "global-config.xml"
        if not gc.exists():
            gc.write_text(common_global_config_xml(use_apikit), encoding="utf-8")

        if raml_info and service_type not in ("PROXY","REC"):
            for recurso, data in sorted(raml_info.items()):
                methods = data.get("methods") or {"retrieve"}
                rname = safe_filename(recurso, "root")
                c_path = client_dir / f"{rname}-client.xml"
                if not c_path.exists():
                    xml = client_file_xml(rname, methods, ctx.get("general_path","/api/*"), use_apikit, raml_cp_value)
                    try: ET.fromstring(xml)
                    except ET.ParseError: xml = f"<mule><flow name='{rname}_client_main'/></mule>"
                    c_path.write_text(xml, encoding="utf-8")
                h_path = handler_dir / f"{rname}-handler.xml"
                if not h_path.exists():
                    xml = handler_file_xml(rname, methods, raml_info)
                    try: ET.fromstring(xml)
                    except ET.ParseError: xml = f"<mule><flow name='{rname}_handler_main'/></mule>"
                    h_path.write_text(xml, encoding="utf-8")
                if len(methods) == 1:
                    o_path = orch_dir / f"{rname}-orchestrator.xml"
                    if not o_path.exists():
                        xml = orchestrator_file_xml(rname, methods, raml_info, single_file=True)
                        try: ET.fromstring(xml)
                        except ET.ParseError: xml = f"<mule><flow name='{rname}_orchestrator_main'/></mule>"
                        o_path.write_text(xml, encoding="utf-8")
                else:
                    for m in sorted(methods):
                        o_path = orch_dir / f"{rname}-{m}-orchestrator.xml"
                        if not o_path.exists():
                            xml = orchestrator_file_xml(rname, {m}, raml_info, single_file=False)
                            try: ET.fromstring(xml)
                            except ET.ParseError: xml = f"<mule><flow name='{rname}_orchestrator_{m}'/></mule>"
                            o_path.write_text(xml, encoding="utf-8")

    enforce_pom_requirements(root, {
        "artifact_id": ctx.get("artifact_id"),
        "project_name": ctx.get("project_name")
    }, use_apikit=False if service_type in ("PROXY","REC") else bool(raml_cp_value))
    write_readme(root, raml_info, ctx, service_type, spec_kind)
    write_validate_script(root)
    if raml_info and service_type not in ("PROXY","REC"):
        write_munit_min(root, raml_info)

    notes = rubric_observaciones(root, use_apikit=False if service_type in ("PROXY","REC") else bool(raml_cp_value), raml_info=raml_info, mode="mule")
    st.session_state.observaciones = notes

    spec_base = Path(st.session_state.get("spec_name") or "proyecto").stem
    out_name = f"{spec_base}.zip"
    out_zip = Path(tempfile.gettempdir()) / out_name
    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as z:
        for p in root.rglob("*"):
            z.write(p, p.relative_to(root))
    return str(out_zip)

def procesar_arquetipo_apigee(arquetipo_dir: str, ctx: dict, spec_bytes: bytes|None, spec_name: str, spec_kind: str):
    """
    Reescribe el arquetipo Reception (Apigee) con metadatos unificados:
    - proxies/default.xml (BasePath)
    - targets/backend.xml (URL)
    - policies/*.xml (Quota/Spike/CORS/VerifyAPIKey, si aplican)
    - resources/jsc/*.js (ajustes menores)
    - resources/oas (coloca OAS del usuario)
    """
    src = Path(arquetipo_dir)
    tmp_dir = Path(tempfile.mkdtemp())
    shutil.copytree(src, tmp_dir, dirs_exist_ok=True)
    roots = [p for p in tmp_dir.rglob("apiproxy") if p.is_dir()]
    root = tmp_dir

    apiproxy_dir = roots[0] if roots else None
    if not apiproxy_dir:
        maybe = list(tmp_dir.rglob("*/apiproxy"))
        if maybe:
            apiproxy_dir = maybe[0]
        else:
            raise RuntimeError("No se encontró carpeta 'apiproxy' en el arquetipo Reception")

    # Nombre del bundle si hay placeholder de carpeta
    bundle_parent = apiproxy_dir.parent
    new_bundle = ctx.get("api_name") or "reference-data"
    try:
        if bundle_parent.name != new_bundle and bundle_parent.parent.exists():
            target_dir = bundle_parent.parent / new_bundle
            if not target_dir.exists():
                shutil.move(str(bundle_parent), str(target_dir))
                apiproxy_dir = target_dir / "apiproxy"
    except Exception:
        pass

    # Colocar OAS
    oas_dir = apiproxy_dir / "resources" / "oas"
    oas_dir.mkdir(parents=True, exist_ok=True)
    oas_target = oas_dir / "rec-reference-data.json"
    if spec_bytes:
        try:
            if spec_kind in ("OAS",) and (spec_name or "").lower().endswith((".yaml",".yml")):
                data_oas = yaml.safe_load(spec_bytes.decode("utf-8","ignore"))
                oas_target.write_text(json.dumps(data_oas, ensure_ascii=False, indent=2), encoding="utf-8")
            else:
                oas_target.write_bytes(spec_bytes)
        except Exception:
            oas_target.write_bytes(spec_bytes)

    # Transformar archivos clave con GPT usando YAML unificado
    files_to_touch = []
    for p in apiproxy_dir.rglob("*"):
        if p.is_file() and (p.suffix.lower() in TEXT_EXTS or p.suffix.lower() in {".xml",".js"}):
            files_to_touch.append(p)

    prog = st.progress(0.0)
    total = len(files_to_touch)
    for i, path in enumerate(files_to_touch, 1):
        prog.progress(i/total)
        try:
            original = path.read_text(encoding="utf-8", errors="ignore")
            nuevo = transformar_archivo_con_gpt(str(path.relative_to(apiproxy_dir)), original, ctx)
            if path.suffix.lower() == ".xml":
                err = validar_xml(nuevo, path.name)
                if err:
                    nuevo = original
            elif path.suffix.lower() in (".yaml",".yml"):
                err = validar_yaml(nuevo, path.name)
                if err:
                    nuevo = original
            path.write_text(nuevo, encoding="utf-8")
        except Exception:
            pass

    # Observaciones Apigee
    notes = rubric_observaciones(root, use_apikit=False, raml_info={}, mode="apigee")
    st.session_state.observaciones = notes

    # ZIP final
    spec_base = Path(spec_name or ctx.get("api_name","reference-data")).stem
    out_name = f"{spec_base}.zip"
    out_zip = Path(tempfile.gettempdir()) / out_name
    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as z:
        for p in root.rglob("*"):
            z.write(p, p.relative_to(root))
    return str(out_zip)

# === NUEVO: ejecución encapsulada de la generación ===
def ejecutar_generacion():
    try:
        if not st.session_state.uploaded_spec:
            st.session_state.messages.append({"role":"assistant","content":"⚠️ Primero adjunta el ZIP de diseño (OpenAPI/RAML incluido dentro)."})
            st.session_state.is_generating = False
            st.session_state.pending_action = None
            return

        # Capa por radio
        choice = st.session_state.get("archetype_choice", "Automático")
        svc = st.session_state.get("service_type", "UNKNOWN")

        # === Leer especificación del usuario y preparar contexto base ===
        raw_ctx = st.session_state.get("ctx_text")
        if not raw_ctx:
            raw_ctx = leer_especificacion(st.session_state.uploaded_spec)

        # === YAML unificado (contexto + reglas de transformación) ===
        yaml_uni = inferir_yaml_unificado(raw_ctx or "", choice if choice != "Automático" else {
            "REC":"Reception", "DOM":"Domain", "BUS":"Business", "PROXY":"Proxy"
        }.get(svc, "Domain"))

        # Adaptación ctx → Mule/Apigee internos para prompts de archivos
        layer = (yaml_uni.get("layer") or "").lower()
        if layer == "reception":
            # Capa REC → Apigee
            st.session_state.rubrics_defs = _load_rubrics_from_root("Rubricas_Scaffold_Apigee.json", "Apigee")
            st.session_state.rubrics_kind = "apigee"
            arquetipo = obtener_arquetipo_reception()
            if not arquetipo:
                st.session_state.messages.append({"role":"assistant","content":"❌ No encontré el arquetipo Reception (Apigee). Copia un *apigee bundle* de referencia en la raíz."})
                st.session_state.is_generating = False
                st.session_state.pending_action = None
                return

            ctx = {
                # Campos que el transformador de archivos espera:
                "layer": "reception",
                "api_display_name": yaml_uni["names"].get("api_display_name"),
                "api_name": yaml_uni["names"].get("api_name"),
                "base_path": yaml_uni["paths"].get("base_path"),
                "target_base_url": yaml_uni["paths"].get("target_base_url"),
                "products": ["auto-generated"],
                "quota": yaml_uni.get("security", {}).get("quota", {"enabled": False}),
                "spike_arrest": yaml_uni.get("security", {}).get("spike_arrest", {"enabled": False}),
                "cors": {"enabled": bool(yaml_uni.get("security", {}).get("cors", True))},
                "auth": {"type": yaml_uni.get("security", {}).get("auth", "none")},
                "transformations": yaml_uni.get("transformations", []),
                "notes": yaml_uni.get("notes","")
            }

            st.session_state.messages.append({"role":"assistant",
                                              "content": f"🧾 YAML unificado:\n```yaml\n{yaml.safe_dump(yaml_uni, sort_keys=False, allow_unicode=True)}\n```"})

            # Bytes reales de especificación
            spec_kind_effective = st.session_state.get("extracted_kind") or st.session_state.get("spec_kind")
            spec_bytes = st.session_state.get("extracted_bytes")
            if not spec_bytes:
                st.session_state.uploaded_spec.seek(0)
                spec_bytes = st.session_state.uploaded_spec.read()

            work_src = preparar_arquetipo_trabajo(arquetipo)

            start_ts = time.time()
            st.session_state.messages.append({"role":"assistant","content":"⚙️ Generando bundle Apigee (Reception): proxies/targets/policies/resources..."})
            salida_zip = procesar_arquetipo_apigee(str(work_src), ctx, spec_bytes, st.session_state.get("spec_name"), spec_kind_effective)
            end_ts = time.time()

            st.session_state.generated_zip = salida_zip
            st.session_state.service_type = "REC"
            label = "RECEPTION"

        else:
            # Capa Mule (Domain/Business/Proxy)
            st.session_state.rubrics_defs = _load_rubrics_from_root("Rubrics_Generation_Mule.json", "Mule")
            st.session_state.rubrics_kind = "mule"

            arquetipo = obtener_arquetipo_generic()
            if not arquetipo:
                st.session_state.messages.append({"role":"assistant","content":"❌ No encontré arquetipo Mule genérico en la raíz."})
                st.session_state.is_generating = False
                st.session_state.pending_action = None
                return

            # Contexto Mule
            general_path = yaml_uni["paths"].get("base_path") or "/api/*"
            # Asegurar wildcard si parece path base
            if general_path and not general_path.endswith("*"):
                if general_path.endswith("/"):
                    general_path = general_path + "*"
                elif general_path.endswith("/*"):
                    pass
                else:
                    # si es tipo "/v1/recurso", usar "/v1/*"
                    segs = general_path.rstrip("/").split("/")
                    general_path = ("/" + segs[1] + "/*") if len(segs) > 1 else "/api/*"

            ctx = {
                "layer": layer,
                "project_name": yaml_uni["names"].get("project_name") or yaml_uni["names"].get("api_display_name","MuleApplication"),
                "artifact_id": yaml_uni["names"].get("artifact_id"),
                "version": yaml_uni["names"].get("version","1.0.0"),
                "group_id": yaml_uni["names"].get("group_id","com.company.domain"),
                "general_path": general_path,
                "base_uri": yaml_uni["paths"].get("base_uri"),
                "upstream_host": yaml_uni["upstream"].get("host"),
                "upstream_protocol": yaml_uni["upstream"].get("protocol"),
                "upstream_path": yaml_uni["upstream"].get("path") or "/",
                "transformations": yaml_uni.get("transformations", []),
                "notes": yaml_uni.get("notes","")
            }

            # bytes de especificación
            spec_kind_effective = st.session_state.get("extracted_kind") or st.session_state.get("spec_kind")
            spec_bytes = st.session_state.get("extracted_bytes")
            if not spec_bytes:
                st.session_state.uploaded_spec.seek(0)
                spec_bytes = st.session_state.uploaded_spec.read()

            work_src = preparar_arquetipo_trabajo(arquetipo)

            start_ts = time.time()
            st.session_state.messages.append({"role":"assistant","content":"⚙️ Reescribiendo arquetipo Mule: POM + XML + README + MUnit + observaciones..."})
            salida_zip = procesar_arquetipo_mule(str(work_src), ctx, spec_bytes, spec_kind_effective)
            end_ts = time.time()

            st.session_state.generated_zip = salida_zip
            st.session_state.service_type = "PROXY" if layer=="proxy" else ("BUS" if layer=="business" else "DOM")
            label = TYPE_LABELS.get(st.session_state.service_type, "UNKNOWN")

            st.session_state.messages.append({"role":"assistant",
                                              "content": f"🧾 YAML unificado:\n```yaml\n{yaml.safe_dump(yaml_uni, sort_keys=False, allow_unicode=True)}\n```"})

        resumen = f"✅ Proyecto generado. Tipo: **{label}**."
        if st.session_state.observaciones:
            resumen += f"\n⚠️ Observaciones: {len(st.session_state.observaciones)} (rúbricas: {st.session_state.rubrics_kind})"
        st.session_state.messages.append({"role":"assistant","content":resumen})
        st.info(f"🔎 Capa seleccionada: **{label}**")

        # 📤 Enviar log a Datadog
        dd_event = {
            "event": "apigee_project_generation" if label=="RECEPTION" else "mule_project_generation",
            "service_type": label,
            "archetype_choice": st.session_state.get("archetype_choice"),
            "spec_name": st.session_state.get("spec_name"),
            "spec_kind": st.session_state.get("spec_kind"),
            "output_zip": Path(st.session_state.generated_zip).name,
            "observations_count": len(st.session_state.get("observaciones", [])),
            "rubrics_count": len(st.session_state.get("rubrics_defs", [])),
            "rubrics_kind": st.session_state.get("rubrics_kind"),
            "yaml_unificado": yaml_uni,
            "chat_history": st.session_state.get("messages", []),
            "started_at": datetime.datetime.utcfromtimestamp(start_ts).isoformat()+"Z",
            "finished_at": datetime.datetime.utcfromtimestamp(end_ts).isoformat()+"Z",
            "duration_seconds": round(end_ts - start_ts, 3),
            "host": socket.gethostname(),
        }
        send_chatlog_to_datadog(dd_event)

    except Exception as e:
        st.session_state.messages.append({"role":"assistant","content":f"⚠️ Generación con advertencias: {e}"})
    finally:
        st.session_state.is_generating = False
        st.session_state.pending_action = None

# ========= Chat/acciones =========

def manejar_mensaje(user_input: str):
    if st.session_state.get("is_generating", False):
        st.session_state.messages.append({
            "role": "assistant",
            "content": "⛔ Generación en curso. Espera a que termine para enviar nuevos mensajes."
        })
        return

    ui = user_input.strip().lower()

    if ui in ("crear proyecto","crea el proyecto","genera el proyecto","crea el proyecto"):
        if not st.session_state.uploaded_spec:
            st.session_state.messages.append({"role":"assistant","content":"⚠️ Primero adjunta el ZIP de diseño con OAS/RAML dentro."})
            return
        st.session_state.is_generating = True
        st.session_state.pending_action = "generate"
        st.session_state.messages.append({"role":"assistant","content":"⏳ Iniciando generación... la entrada quedará deshabilitada hasta finalizar."})
        st.rerun()
        return
    else:
        st.session_state.messages.append({"role":"assistant","content":"💬 Escribe \"crea el proyecto\" para generar el ZIP del proyecto a partir de tu diseño."})

# ========= Upload =========

spec = st.file_uploader(
    "Adjunta el ZIP de diseño (mx-api-bc-dom-* o mx-api-bc-rec-*). Debe contener openapi.* o .raml",
    type=["zip"]
)
if spec and st.session_state.uploaded_spec is None:
    st.session_state.uploaded_spec = spec
    st.session_state.spec_name = spec.name
    stype = _map_prefix_to_type(spec.name)
    st.session_state.service_type = stype if stype else "UNKNOWN"
    # Pre-parse: deja ctx_text y extracted_* listos
    leer_especificacion(spec)
    st.session_state.messages.append({
        "role": "assistant",
        "content": f"📦 Especificación \"{spec.name}\" cargada. Elige la capa y escribe \"crea el proyecto\"."
    })

# === Selector de capa/arquetipo (Domain / Reception / Business / Proxy) ===
if st.session_state.get("uploaded_spec"):
    choices = ["Domain", "Reception", "Business", "Proxy", "Automático"]
    default_idx = 4
    # Si el ZIP sugiere tipo, marcamos:
    inferred = st.session_state.get("service_type")
    if inferred in ("REC","DOM","BUS","PROXY"):
        default_idx = {"DOM":0, "REC":1, "BUS":2, "PROXY":3}[inferred]
    st.session_state.archetype_choice = st.radio(
        "Selecciona la capa para generar el proyecto",
        choices,
        index=default_idx,
        horizontal=True,
        key="arch_choice_radio"
    )

# ====== Render historial ======
with st.container():
    for msg in st.session_state.messages:
        avatar = user_avatar if msg["role"]=="user" else assistant_avatar
        bubble = "user-bubble" if msg["role"]=="user" else "assistant-bubble"
        who = "user-message" if msg["role"]=="user" else "assistant-message"
        st.markdown(
            f'<div class="chat-message {who}"><img src="{avatar}" class="avatar"><div class="message-bubble {bubble}">{msg["content"]}</div></div>',
            unsafe_allow_html=True
        )

# ====== Observaciones ======
if st.session_state.get("observaciones"):
    st.markdown("### ⚠️ Observaciones de Calidad (con rúbricas)")
    for o in st.session_state.observaciones:
        st.markdown(f"- {o}", unsafe_allow_html=True)

# === Oculta chat_input si hay generación ===
if st.session_state.get("is_generating", False):
    st.markdown("<style>div[data-testid='stChatInput']{display:none !important}</style>", unsafe_allow_html=True)

# === Si hay generación pendiente, ejecutarla ===
if st.session_state.is_generating and st.session_state.pending_action == "generate":
    with st.spinner("Generando proyecto..."):
        ejecutar_generacion()
    st.rerun()

# ====== Entrada chat ======
chat_area = st.empty()
if st.session_state.get("is_generating", False):
    with chat_area:
        st.markdown(
            "<div style='opacity:.75; padding:10px; border:1px dashed #bbb; border-radius:10px;'>"
            "⏳ <b>Generando proyecto…</b> El chat está temporalmente deshabilitado."
            "</div>",
            unsafe_allow_html=True
        )
else:
    with chat_area:
        user_input = st.chat_input("Escribe aquí...", key="chat_input_main_v3")
        if user_input:
            st.session_state.messages.append({"role":"user","content":user_input})
            manejar_mensaje(user_input)
            st.rerun()

# ====== Descarga ======
if st.session_state.generated_zip:
    tipo = st.session_state.get("service_type","UNKNOWN")
    label = TYPE_LABELS.get(tipo, tipo)
    st.info(f"🔎 Capa: **{label}**")
    with open(st.session_state.generated_zip, "rb") as f:
        st.download_button("⬇️ Descargar Proyecto (.zip)", f,
                           Path(st.session_state.generated_zip).name,
                           "application/zip")
